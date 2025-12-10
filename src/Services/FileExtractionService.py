"""
Service pour l'extraction de texte depuis différents formats de fichiers
"""
import io
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document
from src.Utils.BaseError import BaseError


class FileExtractionService:
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extrait le texte d'un fichier PDF
        
        Args:
            file_content: Contenu binaire du fichier PDF
            
        Returns:
            Texte extrait du PDF
            
        Raises:
            BaseError: Si l'extraction échoue
        """
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            full_text = '\n'.join(text_parts).strip()
            
            if not full_text:
                raise BaseError("Le PDF ne contient pas de texte extractible (peut-être une image scannée)", 400)
            
            return full_text
        except Exception as e:
            if isinstance(e, BaseError):
                raise
            raise BaseError(f"Erreur lors de l'extraction du texte du PDF: {str(e)}", 500)
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extrait le texte d'un fichier DOCX
        
        Args:
            file_content: Contenu binaire du fichier DOCX
            
        Returns:
            Texte extrait du DOCX
            
        Raises:
            BaseError: Si l'extraction échoue
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extraire aussi le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_parts).strip()
            
            if not full_text:
                raise BaseError("Le fichier DOCX ne contient pas de texte extractible", 400)
            
            return full_text
        except Exception as e:
            if isinstance(e, BaseError):
                raise
            raise BaseError(f"Erreur lors de l'extraction du texte du DOCX: {str(e)}", 500)
    
    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        """
        Extrait le texte d'un fichier texte
        
        Args:
            file_content: Contenu binaire du fichier texte
            
        Returns:
            Texte extrait
            
        Raises:
            BaseError: Si l'extraction échoue
        """
        try:
            # Essayer UTF-8 d'abord
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Essayer latin-1 si UTF-8 échoue
                text = file_content.decode('latin-1')
            
            text = text.strip()
            
            if not text:
                raise BaseError("Le fichier texte est vide", 400)
            
            return text
        except Exception as e:
            if isinstance(e, BaseError):
                raise
            raise BaseError(f"Erreur lors de la lecture du fichier texte: {str(e)}", 500)
    
    @staticmethod
    def extract_text_from_file(file_content: bytes, file_extension: str) -> str:
        """
        Extrait le texte d'un fichier selon son extension
        
        Args:
            file_content: Contenu binaire du fichier
            file_extension: Extension du fichier (.pdf, .docx, .txt)
            
        Returns:
            Texte extrait
            
        Raises:
            BaseError: Si le format n'est pas supporté ou si l'extraction échoue
        """
        extension = file_extension.lower().lstrip('.')
        
        if extension == 'pdf':
            return FileExtractionService.extract_text_from_pdf(file_content)
        elif extension == 'docx':
            return FileExtractionService.extract_text_from_docx(file_content)
        elif extension == 'txt':
            return FileExtractionService.extract_text_from_txt(file_content)
        else:
            raise BaseError(
                f"Format de fichier non supporté: {extension}. Formats acceptés: PDF, DOCX, TXT",
                400
            )

